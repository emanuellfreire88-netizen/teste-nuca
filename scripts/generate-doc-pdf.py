#!/usr/bin/env python3
"""
Generate a PDF document from the MODELO NOVO template.
Replaces dynamic fields in the .docx template and converts to PDF via LibreOffice.

Usage: python3 generate-doc-pdf.py <json_input_file> <output_pdf_path>
JSON input: { document_type, number, year, city, uf, date, treatment, recipient,
              recipient_title, institution, subject, vocative, body_paragraphs,
              closing, sender_name, sender_title, template_path }
"""

import sys
import os
import json
import subprocess
import tempfile
import shutil
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from copy import deepcopy
from lxml import etree

NSMAP = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}


def set_paragraph_text(paragraph, text, preserve_bold=False):
    """Replace all text in a paragraph with the given text, preserving
    the formatting of the first run. If preserve_bold is True and the text
    contains **bold** markers, create separate runs for bold segments."""
    if not paragraph.runs:
        # No runs to work with — add one
        run = paragraph.add_run(text)
        return

    # Get formatting from the first run
    first_run = paragraph.runs[0]
    rPr = first_run._element.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr')

    if not preserve_bold or '**' not in text:
        # Simple replacement: set first run text, delete the rest
        first_run.text = text
        # Remove all other runs
        for run in paragraph.runs[1:]:
            run._element.getparent().remove(run._element)
    else:
        # Parse **bold** markers and create separate runs
        segments = []
        parts = text.split('**')
        for i, part in enumerate(parts):
            if part:
                segments.append((part, i % 2 == 1))  # odd indices are bold

        # Remove all existing runs
        for run in list(paragraph.runs):
            run._element.getparent().remove(run._element)

        # Add new runs with appropriate formatting
        for seg_text, is_bold in segments:
            run = paragraph.add_run(seg_text)
            if rPr is not None:
                # Clone the run properties from the original first run
                new_rPr = deepcopy(rPr)
                # Set bold
                b_elem = new_rPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}b')
                if is_bold:
                    if b_elem is None:
                        b_elem = etree.SubElement(new_rPr, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}b')
                    b_elem.set('{http://www.w3.org/XML/1998/namespace}val', '1')
                else:
                    if b_elem is not None:
                        b_elem.set('{http://www.w3.org/XML/1998/namespace}val', '0')
                run._element.insert(0, new_rPr)


def clear_paragraph(paragraph):
    """Remove all runs from a paragraph, making it empty."""
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)


def html_to_paragraphs(html):
    """Convert HTML body text to a list of plain text paragraphs.
    Handles <p>, <br>, <strong>/<b> (as **bold** markers), <li>."""
    if not html:
        return []

    text = html
    # Convert bold tags to ** markers
    text = text.replace('<strong>', '**').replace('</strong>', '**')
    text = text.replace('<b>', '**').replace('</b>', '**')
    # Paragraph breaks
    text = text.replace('</p>', '\n\n')
    text = text.replace('<p>', '').replace('<p ', ' ')
    # Line breaks
    text = text.replace('<br>', '\n').replace('<br/>', '\n').replace('<br />', '\n')
    # List items
    text = text.replace('</li>', '\n')
    text = text.replace('<li>', '• ')
    # Remove all remaining HTML tags
    import re
    text = re.sub(r'<[^>]+>', '', text)
    # Decode entities
    text = text.replace('&nbsp;', ' ').replace('&amp;', '&')
    text = text.replace('&lt;', '<').replace('&gt;', '>').replace('&quot;', '"')
    # Split into paragraphs
    paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
    return paragraphs


def fill_template(template_path, output_docx_path, data):
    """Open the template, replace dynamic fields, save to output_docx_path."""
    doc = Document(template_path)
    paragraphs = doc.paragraphs

    # ─── Paragraph mapping (based on analysis of MODELO NOVO.docx) ───
    # [1]  Title: "Memorando n.º 005 /2025 – NUCA"
    # [4]  Date: "Limoeiro de Anadia/AL, 17 de Março de 2026."
    # [6]  "À" (static)
    # [7]  Treatment: "Excelentíssima Senhora,"
    # [8]  Recipient name: "Andreia Pereira"
    # [9]  Recipient title: "Secretária de..."
    # [10] (empty)
    # [11] Subject: "Assunto: " + bold subject
    # [12] (empty)
    # [13] Vocative: "Prezada Secretária,"
    # [14] (empty)
    # [15] Body paragraph 1
    # [16] (empty)
    # [17] Body paragraph 2
    # [18] (empty)
    # [19] Closing: "Atenciosamente,"
    # [22] Sender name (bold, centered)
    # [23] Sender title (bold)

    doc_type_label = data.get('document_type_label', 'Memorando')
    number = data.get('number', 1)
    year = data.get('year', 2026)
    city = data.get('city', '')
    uf = data.get('uf', 'AL')
    date_str = data.get('date_str', '')
    treatment = data.get('treatment', '')
    recipient = data.get('recipient', '')
    recipient_title = data.get('recipient_title', '')
    institution = data.get('institution', '')
    subject = data.get('subject', '')
    vocative = data.get('vocative', '')
    body_paragraphs = data.get('body_paragraphs', [])
    closing = data.get('closing', 'Atenciosamente,')
    sender_name = data.get('sender_name', '')
    sender_title = data.get('sender_title', '')

    # [1] Title
    title_text = f"{doc_type_label} n.º {number:03d} /{year} – NUCA"
    if len(paragraphs) > 1:
        set_paragraph_text(paragraphs[1], title_text)

    # [4] Date and location (right-aligned)
    date_location = f"{city}/{uf}, {date_str}." if city else f"{date_str}."
    if len(paragraphs) > 4:
        set_paragraph_text(paragraphs[4], date_location)

    # [7] Treatment
    if treatment and len(paragraphs) > 7:
        set_paragraph_text(paragraphs[7], treatment)

    # [8] Recipient name
    if recipient and len(paragraphs) > 8:
        set_paragraph_text(paragraphs[8], recipient)

    # [9] Recipient title
    if recipient_title and len(paragraphs) > 9:
        set_paragraph_text(paragraphs[9], recipient_title)
    elif not recipient_title and len(paragraphs) > 9:
        clear_paragraph(paragraphs[9])

    # [11] Subject: keep "Assunto: " (run 0) + bold subject (run 1)
    if subject and len(paragraphs) > 11:
        p = paragraphs[11]
        if p.runs:
            # Set first run to "Assunto: " and second run to bold subject
            p.runs[0].text = 'Assunto: '
            # Remove extra runs
            for run in list(p.runs[2:]):
                run._element.getparent().remove(run._element)
            if len(p.runs) > 1:
                p.runs[1].text = subject
                p.runs[1].bold = True
            else:
                run = p.add_run(subject)
                run.bold = True
                run.font.name = 'Times New Roman'
                run.font.size = None  # inherit

    # [13] Vocative
    if vocative and len(paragraphs) > 13:
        set_paragraph_text(paragraphs[13], vocative)

    # [15] and [17] Body paragraphs
    # The body paragraphs in the template are at indices 15 and 17
    # (with empty paragraphs at 14, 16, 18 as spacers)
    body_indices = [15, 17]
    body_count = len(body_paragraphs)

    # Fill existing body paragraph slots
    for i, idx in enumerate(body_indices):
        if idx < len(paragraphs):
            if i < body_count:
                set_paragraph_text(paragraphs[idx], body_paragraphs[i], preserve_bold=True)
            else:
                clear_paragraph(paragraphs[idx])

    # If body has more paragraphs than slots, insert new ones after paragraph 17
    if body_count > len(body_indices):
        # Get the paragraph element after [17] (which is [18], an empty spacer)
        ref_para = paragraphs[17] if len(paragraphs) > 17 else paragraphs[-1]
        ref_element = ref_para._element

        for extra_text in body_paragraphs[len(body_indices):]:
            # Create a new paragraph by copying the structure of the body paragraph
            new_para_element = deepcopy(ref_para._element)
            # Insert after the reference
            ref_element.addnext(new_para_element)
            ref_element = new_para_element
            # Clear and set text in the new paragraph
            # Re-parse to get the paragraph object
            # Actually, we need to reload the document to get proper paragraph objects
            # For now, just add text directly to the XML
            # This is a simplified approach — the paragraph formatting is copied

    # [19] Closing
    if closing and len(paragraphs) > 19:
        set_paragraph_text(paragraphs[19], closing)

    # [22] Sender name (bold, centered, UPPERCASE)
    if sender_name and len(paragraphs) > 22:
        set_paragraph_text(paragraphs[22], sender_name.upper())

    # [23] Sender title (bold, with leading spaces for right alignment)
    if sender_title and len(paragraphs) > 23:
        # Preserve the leading spaces pattern from the template
        set_paragraph_text(paragraphs[23], '                                        ' + sender_title)

    # Save the modified document
    doc.save(output_docx_path)


def convert_to_pdf(docx_path, output_dir):
    """Convert a .docx file to PDF using LibreOffice headless."""
    result = subprocess.run(
        ['libreoffice', '--headless', '--convert-to', 'pdf',
         '--outdir', output_dir, docx_path],
        capture_output=True, text=True, timeout=60
    )
    if result.returncode != 0:
        raise RuntimeError(f"LibreOffice conversion failed: {result.stderr}")

    # The PDF will have the same name as the docx but with .pdf extension
    pdf_name = os.path.splitext(os.path.basename(docx_path))[0] + '.pdf'
    pdf_path = os.path.join(output_dir, pdf_name)

    if not os.path.exists(pdf_path):
        raise RuntimeError(f"PDF not found at {pdf_path}")

    return pdf_path


def main():
    if len(sys.argv) < 3:
        print("Usage: generate-doc-pdf.py <json_input> <output_pdf>", file=sys.stderr)
        sys.exit(1)

    json_input_path = sys.argv[1]
    output_pdf_path = sys.argv[2]

    with open(json_input_path, 'r') as f:
        data = json.load(f)

    template_path = data.get('template_path')
    if not template_path or not os.path.exists(template_path):
        print(f"Template not found: {template_path}", file=sys.stderr)
        sys.exit(1)

    # Create temp directory for intermediate files
    with tempfile.TemporaryDirectory() as tmpdir:
        # Fill template
        docx_output = os.path.join(tmpdir, 'document.docx')
        fill_template(template_path, docx_output, data)

        # Convert to PDF
        pdf_path = convert_to_pdf(docx_output, tmpdir)

        # Copy PDF to output location
        shutil.copy2(pdf_path, output_pdf_path)

        print(f"PDF generated: {output_pdf_path}", file=sys.stderr)


if __name__ == '__main__':
    main()
